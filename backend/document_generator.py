from docx import Document
from docx.shared import Pt , Inches
from datetime import date
import os

# ── Template folder path ──────────────────────────────────────────────────────
TEMPLATES_DIR = "documents/templates"
OUTPUT_DIR    = "files"

# Make sure output folder exists
os.makedirs(OUTPUT_DIR, exist_ok=True)

# ── Helper: replace all placeholders in a document ───────────────────────────
def fill_template(template_name: str, placeholders: dict, 
                  output_filename: str, city: str = None) -> str:
    
    template_path = os.path.join(TEMPLATES_DIR, template_name)

    if not os.path.exists(template_path):
       raise FileNotFoundError(f"Template not found: {template_path}")
    
    doc = Document(template_path)
    
    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        for key, value in placeholders.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, str(value))
    
    # Replace placeholders in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in placeholders.items():
                        if key in paragraph.text:
                            for run in paragraph.runs:
                                if key in run.text:
                                    run.text = run.text.replace(key, str(value))
    
    # ── Add signature based on city ───────────────────────────────────────────
    if city:
        add_signature(doc, city)
    
    output_path = os.path.join(OUTPUT_DIR, output_filename)
    doc.save(output_path)
    return output_path

# ── Number to French words (for salary) ──────────────────────────────────────
def number_to_french(n: float) -> str:
    units = ['', 'un', 'deux', 'trois', 'quatre', 'cinq', 'six', 'sept', 'huit', 'neuf',
             'dix', 'onze', 'douze', 'treize', 'quatorze', 'quinze', 'seize',
             'dix-sept', 'dix-huit', 'dix-neuf']
    tens  = ['', '', 'vingt', 'trente', 'quarante', 'cinquante', 'soixante',
             'soixante', 'quatre-vingt', 'quatre-vingt']
    
    n = int(n)
    if n == 0: return 'zéro'
    if n < 0:  return 'moins ' + number_to_french(-n)
    
    result = ''
    if n >= 1000:
        result += number_to_french(n // 1000) + ' mille '
        n %= 1000
    if n >= 100:
        if n // 100 == 1:
            result += 'cent '
        else:
            result += units[n // 100] + ' cent '
        n %= 100
    if n >= 20:
        t = n // 10
        u = n % 10
        if t == 7 or t == 9:
            result += tens[t] + '-' + units[10 + u] + ' '
        elif u == 1 and t != 8:
            result += tens[t] + ' et un '
        elif u > 0:
            result += tens[t] + '-' + units[u] + ' '
        else:
            result += tens[t] + ' '
    elif n > 0:
        result += units[n] + ' '
    
    return result.strip() + ' dirhams'

# ── Format date in French ─────────────────────────────────────────────────────
def format_date_french(d) -> str:
    months = ['janvier','février','mars','avril','mai','juin',
              'juillet','août','septembre','octobre','novembre','décembre']
    if isinstance(d, str):
        from datetime import datetime
        d = datetime.strptime(d, '%Y-%m-%d').date()
    return f"{d.day} {months[d.month - 1]} {d.year}"

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT GENERATORS
# ══════════════════════════════════════════════════════════════════════════════

def generate_attestation_travail(employee: dict) -> str:
    today = date.today()
    placeholders = {
        "{{nom_complet}}":    f"{employee['first_name']} {employee['last_name']}",
        "{{genre}}":          employee['gender'],
        "{{date_naissance}}": format_date_french(employee['birth_date']),
        "{{departement}}":    employee['department'],
        "{{poste}}":          employee['position'],
        "{{type_contrat}}":   employee['contract_type'],
        "{{date_embauche}}":  format_date_french(employee['hire_date']),
        "{{ville}}":          employee['city'],
        "{{date_generation}}": format_date_french(today),
    }
    last_name = employee['last_name'].replace(' ', '_')
    filename = f"attestation_travail_{last_name}_{employee['first_name']}_{today}.docx"
    # FIXED: Added city parameter here so the signature triggers properly!
    return fill_template("attestation_travail.docx", placeholders, filename, city=employee['city'])


def generate_attestation_salaire(employee: dict, objet: str = "usage personnel") -> str:
    today = date.today()
    placeholders = {
        "{{nom_complet}}":     f"{employee['first_name']} {employee['last_name']}",
        "{{departement}}":     employee['department'],
        "{{poste}}":           employee['position'],
        "{{type_contrat}}":    employee['contract_type'],
        "{{date_embauche}}":   format_date_french(employee['hire_date']),
        "{{salaire}}":         f"{employee['salary']:,.2f}",
        "{{salaire_lettres}}": number_to_french(employee['salary']),
        "{{objet}}":           objet,
        "{{ville}}":           employee['city'],
        "{{date_generation}}": format_date_french(today),
    }
    last_name = employee['last_name'].replace(' ', '_')
    filename = f"attestation_salaire_{last_name}_{employee['first_name']}_{today}.docx"
    return fill_template("attestation_salaire.docx", placeholders, filename, city=employee['city'])


def generate_lettre_conge(employee: dict, leave_request: dict) -> str:
    today = date.today()
    placeholders = {
        "{{nom_complet}}":     f"{employee['first_name']} {employee['last_name']}",
        "{{departement}}":     employee['department'],
        "{{poste}}":           employee['position'],
        "{{type_conge}}":      leave_request['leave_type'],
        "{{date_debut}}":      format_date_french(leave_request['start_date']),
        "{{date_fin}}":        format_date_french(leave_request['end_date']),
        "{{duree}}":           str(leave_request['duration_days']),
        "{{solde_restant}}":   str(employee['leave_balance_days']),
        "{{ville}}":           employee['city'],
        "{{date_generation}}": format_date_french(today),
    }
    last_name = employee['last_name'].replace(' ', '_')
    filename = f"lettre_conge_{last_name}_{employee['first_name']}_{today}.docx"
    return fill_template("lettre_conge.docx", placeholders, filename, city=employee['city'])


def generate_bulletin_paie(employee: dict, month: str, year: int) -> str:
    today = date.today()
    
    # Simple salary calculations
    salaire_brut   = float(employee['salary'])
    cnss_employee  = round(salaire_brut * 0.0448, 2)
    amo_employee   = round(salaire_brut * 0.0226, 2)  
    ir_base        = salaire_brut - cnss_employee - amo_employee
    ir             = round(ir_base * 0.15, 2)
    salaire_net    = round(salaire_brut - cnss_employee - amo_employee - ir, 2)
    
    placeholders = {
        "{{nom_complet}}":     f"{employee['first_name']} {employee['last_name']}",
        "{{departement}}":     employee['department'],
        "{{poste}}":           employee['position'],
        "{{type_contrat}}":    employee['contract_type'],
        "{{date_embauche}}":   format_date_french(employee['hire_date']),
        "{{mois}}":            month,
        "{{annee}}":           str(year),
        "{{salaire_brut}}":    f"{salaire_brut:,.2f}",
        "{{cnss}}":            f"{cnss_employee:,.2f}",
        "{{amo}}":             f"{amo_employee:,.2f}",
        "{{ir}}":              f"{ir:,.2f}",
        "{{salaire_net}}":     f"{salaire_net:,.2f}",
        "{{ville}}":           employee['city'],
        "{{date_generation}}": format_date_french(today),
    }
    last_name = employee['last_name'].replace(' ', '_')
    filename = f"bulletin_paie_{last_name}_{employee['first_name']}_{today}.docx"
    return fill_template("bulletin_paie.docx", placeholders, filename, city=employee['city'])


def generate_certificat_travail(employee: dict, end_date: str, reason: str = "fin de contrat") -> str:
    today = date.today()
    placeholders = {
        "{{nom_complet}}":     f"{employee['first_name']} {employee['last_name']}",
        "{{genre}}":           employee['gender'],
        "{{departement}}":     employee['department'],
        "{{poste}}":           employee['position'],
        "{{type_contrat}}":    employee['contract_type'],
        "{{date_embauche}}":   format_date_french(employee['hire_date']),
        "{{date_fin}}":        format_date_french(end_date),
        "{{motif}}":           reason,
        "{{ville}}":           employee['city'],
        "{{date_generation}}": format_date_french(today),
    }
    last_name = employee['last_name'].replace(' ', '_')
    filename = f"certificat_travail_{last_name}_{employee['first_name']}_{today}.docx"
    return fill_template("certificat_travail.docx", placeholders, filename, city=employee['city'])


def add_signature(doc, city: str):
    """Bulletproof safety check to prevent document corruption"""
    print("\n" + "🚨" * 20)
    print(f"[SIGNATURE DEBUG] Raw city received from database: '{city}'")
    
    if not city:
        print("[SIGNATURE DEBUG] ERROR: City field is None or empty!")
        print("🚨" * 20 + "\n")
        return
        
    clean_city = str(city).strip().lower()
    print(f"[SIGNATURE DEBUG] Cleaned city name being looked up: '{clean_city}'")

    current_file_dir = os.path.dirname(os.path.abspath(__file__))
    base_dir = current_file_dir
    while base_dir and not os.path.exists(os.path.join(base_dir, "documents")) and base_dir != os.path.dirname(base_dir):
        base_dir = os.path.dirname(base_dir)
        
    if not os.path.exists(os.path.join(base_dir, "documents")):
        base_dir = os.path.abspath(os.path.join(current_file_dir, ".."))

    signatures = {
        "casablanca": os.path.join(base_dir, "documents", "signatures", "signature_casablanca.png"),
        "rabat":      os.path.join(base_dir, "documents", "signatures", "signature_rabat.png"),
        "tanger":     os.path.join(base_dir, "documents", "signatures", "signature_tanger.png"),
        "tangier":    os.path.join(base_dir, "documents", "signatures", "signature_tanger.png"),
    }
    
    sig_path = signatures.get(clean_city)
    print(f"[SIGNATURE DEBUG] Looking for file at absolute path:\n -> {sig_path}")
    
    file_exists = os.path.exists(sig_path) if sig_path else False
    print(f"[SIGNATURE DEBUG] Does file exist on disk? {file_exists}")
    print("🚨" * 20 + "\n")

    if sig_path and file_exists:
        for paragraph in doc.paragraphs:
            if "[Signature et Cachet]" in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture(sig_path, width=Inches(1.6))
                print("[SIGNATURE DEBUG] SUCCESS: Image added successfully.")
                break
    else:
        print(f"[SIGNATURE DEBUG] CRITICAL CRASH PREVENTED: Skipped adding image for '{city}' because the file path wasn't found.")